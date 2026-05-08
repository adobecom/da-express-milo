"""
Milo doc builder — reusable helpers for producing AEM/DA-compatible .docx files.

Each block in a Milo doc is a bordered table with a merged gray header row
showing the block name, then 2-column content rows. Sections are separated by a
horizontal "section break" marker.

Typical feature authoring flow:
    from .claude.tools.build_milo_doc import (
        Document, add_block, add_section_break, add_runs,
    )
    from docx.shared import Cm, Pt

    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(1.5)

    add_block(doc, 'frictionless-quick-action', [
        [HEADER_CELL_CONTENT],                       # merged row spanning cols
        [LEFT_CELL, RIGHT_CELL],                     # 2-col row
        [[('p', [('text', 'Quick-Action')])], [('p', [('text', 'edit-video')])]],
    ], col_widths=[3.3, 3.3])

    add_block(doc, 'section-metadata',
              [[[('p', [('text', 'showwith')])], [('p', [('text', 'fqa-qualified-desktop')])]]],
              col_widths=[3.3, 3.3])

    add_section_break(doc)
    doc.save('page.docx')

Dependencies (install via pip, not brew):
    pip install python-docx requests
"""
import io
import requests
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADER_FILL = "EFEFEF"
BORDER_COLOR = "D0D0D0"
LINK_COLOR = "1473E6"

# Common Adobe/Milo URLs — re-exported so feature build.py scripts don't redefine them.
TERMS_URL = "https://www.adobe.com/legal/terms.html"
PRIVACY_URL = "https://www.adobe.com/privacy/policy.html"

# ----- low-level helpers ---------------------------------------------------

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_borders(cell, color=BORDER_COLOR, sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_table_borders(table, color=BORDER_COLOR, sz="4"):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def merge_row_cells(table, row_idx):
    """Merge all cells in a row into one."""
    row = table.rows[row_idx]
    first = row.cells[0]
    for c in row.cells[1:]:
        first = first.merge(c)
    return first


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), LINK_COLOR)
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_runs(paragraph, parts):
    """parts = list of tuples:
        ('text', str)
        ('em', str)           — italic
        ('link', text, url)
        ('br',)               — line break
    """
    for p in parts:
        kind = p[0]
        if kind == 'text':
            paragraph.add_run(p[1])
        elif kind == 'em':
            r = paragraph.add_run(p[1])
            r.italic = True
        elif kind == 'link':
            add_hyperlink(paragraph, p[1], p[2])
        elif kind == 'br':
            paragraph.add_run().add_break()


def _apply_heading_style(paragraph, level):
    """Apply Word's built-in `Heading N` paragraph style so DA/EDS converts the
    paragraph to a real <hN> element on ingest.

    A bold run at a custom font size is NOT sufficient — DA only emits <h1>/<h2>/<h3>
    when the paragraph has `<w:pStyle w:val="HeadingN"/>` in its properties. Every
    Document Word opens has Heading 1–9 as built-in styles by default.
    """
    try:
        paragraph.style = paragraph.part.document.styles[f'Heading {level}']
    except KeyError:
        # Style not registered on this document — fall back to bold at approximate size.
        run = paragraph.runs[-1] if paragraph.runs else paragraph.add_run()
        run.bold = True
        run.font.size = Pt({1: 20, 2: 16, 3: 13}.get(level, 13))


def write_cell(cell, content_blocks, bold_all=False):
    """content_blocks: list of blocks. Each block is one of:
        ('p', parts)            — paragraph (parts passed to add_runs)
        ('h', level, text)      — heading. Applies Word's `Heading N` paragraph style
                                  so DA/EDS ingest converts it to a real <hN>. Runs
                                  inherit font size/weight from the style — do NOT
                                  also set bold or font size manually.
        ('ul', [items])         — bulleted list (each item is a parts list)
        ('img', url, alt)       — image fetched from url; falls back to [image: alt] on error
    """
    cell.text = ""
    first = True
    for block in content_blocks:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        kind = block[0]
        if kind == 'p':
            parts = block[1]
            add_runs(p, parts)
            for run in p.runs:
                if bold_all:
                    run.bold = True
        elif kind == 'h':
            level, text = block[1], block[2]
            p.add_run(text)
            _apply_heading_style(p, level)
        elif kind == 'ul':
            items = block[1]
            for i, item_parts in enumerate(items):
                lp = p if i == 0 else cell.add_paragraph()
                lp.style = cell.part.document.styles['List Bullet']
                add_runs(lp, item_parts)
        elif kind == 'img':
            url, alt = block[1], block[2]
            try:
                data = requests.get(url, timeout=20).content
                stream = io.BytesIO(data)
                run = p.add_run()
                run.add_picture(stream, width=Inches(2.6))
            except Exception:
                r = p.add_run(f"[image: {alt}]")
                r.italic = True


# ----- block builders ------------------------------------------------------

def add_block(doc, name, rows, col_widths=None):
    """
    Write a Milo block as a bordered table with a gray header row.

    name        — block name shown in gray header row (e.g. 'columns (fullsize)').
                  Variants go in parens and become CSS modifier classes.
    rows        — list of rows. Each row is a list of cells. Each cell is a
                  list of content_blocks (see write_cell). A row with a single
                  cell is rendered full-width (merged across all columns).
    col_widths  — optional list of column widths in inches.
    """
    ncols = max((len(r) for r in rows), default=1)
    ncols = max(ncols, 1)

    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.autofit = False
    set_table_borders(table)

    hcell = merge_row_cells(table, 0)
    set_cell_shading(hcell, HEADER_FILL)
    hcell.text = ""
    hp = hcell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(name)
    hr.bold = True
    hr.font.size = Pt(11)

    for ri, row in enumerate(rows, start=1):
        if len(row) == 1 and ncols > 1:
            c = merge_row_cells(table, ri)
            write_cell(c, row[0])
        else:
            for ci, cell_blocks in enumerate(row):
                write_cell(table.rows[ri].cells[ci], cell_blocks)

    if col_widths and len(col_widths) == ncols:
        for ri in range(len(table.rows)):
            for ci, w in enumerate(col_widths):
                try:
                    table.rows[ri].cells[ci].width = Inches(w)
                except Exception:
                    pass

    doc.add_paragraph()
    return table


def add_section_break(doc):
    """Insert a Milo/DA-compatible section break between blocks.

    Emits TWO signals so the converter picks up on whichever it supports:
      1. A paragraph containing `---` — matches Milo's canonical markdown
         section-boundary convention. DA's docx-to-markdown importer
         translates this to a section break reliably.
      2. A native Word continuous section break (`<w:sectPr type="continuous">`)
         via `doc.add_section()` — Word renders the thin-line "SECTION BREAK"
         indicator natively, and some importers recognize this as the semantic
         boundary too.

    Using just (2) alone broke section detection on DA in practice (all blocks
    merged into one section → fallback hero rendered alongside qualified
    sections, and its plain-link CTA produced a redirect on click). The `---`
    paragraph is the load-bearing signal; the native break is belt-and-braces.
    """
    p = doc.add_paragraph("---")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_section(WD_SECTION.CONTINUOUS)


# ----- high-level block helpers (schema-driven, content-only API) -----------
# Each of these encapsulates a Milo authoring pattern so feature build.py
# scripts only pass content — never row shape, column widths, or merged cells.


def add_h2(doc, text):
    """Standalone Heading 2 paragraph between blocks.

    Applies Word's built-in `Heading 2` style so DA ingest emits a real <h2>.
    Use this for section headings that sit OUTSIDE a block's table (e.g. the
    "How to compress a JPEG." heading above the how-to steps block).
    """
    p = doc.add_paragraph(text, style='Heading 2')
    return p


def add_showwith(doc, value, col_widths=(3.3, 3.3)):
    """section-metadata block with a single `showwith = value` row."""
    return add_block(
        doc, "section-metadata",
        [[[('p', [('text', 'showwith')])], [('p', [('text', value)])]]],
        col_widths=list(col_widths),
    )


def add_section_metadata(doc, pairs, col_widths=(3.3, 3.3)):
    """Generic section-metadata block. pairs = dict of key → value strings."""
    rows = [
        [[('p', [('text', k)])], [('p', [('text', v)])]]
        for k, v in pairs.items()
    ]
    return add_block(doc, "section-metadata", rows, col_widths=list(col_widths))


def add_columns_fullsize_hero(
    doc, *, headline, subhead, cta_text, cta_url,
    upload_animation_url, col_widths=(3.3, 3.3),
):
    """Fallback hero for non-qualified browsers (columns fullsize variant)."""
    return add_block(
        doc, "columns (fullsize)",
        [[
            [
                ('h', 1, headline),
                ('p', [('text', subhead)]),
                ('p', [('link', cta_text, cta_url)]),
            ],
            [
                ('p', [('link', "Upload animation (MP4)", upload_animation_url)]),
            ],
        ]],
        col_widths=list(col_widths),
    )


def add_frictionless_quick_action(
    doc, *, headline, subhead, upload_animation_url,
    upload_heading_text, upload_heading_em,
    upload_cta_text, upload_cta_url,
    file_restrictions_text, quick_action_id,
    terms_url=TERMS_URL, privacy_url=PRIVACY_URL,
    col_widths=(3.3, 3.3),
):
    """Desktop frictionless-quick-action hero block.

    Three-row pattern:
      (headline + subhead, empty)
      (video link, upload card with heading + CTA + restrictions + ToS)
      (Quick-Action, <id>)

    The upload heading renders as "<upload_heading_text>\\nor <upload_heading_em>"
    with the em part italicised.
    """
    return add_block(
        doc, "frictionless-quick-action",
        [
            [
                [('h', 1, headline), ('p', [('text', subhead)])],
                [('p', [('text', '')])],
            ],
            [
                [('p', [('link', "Alternate video source (MP4)", upload_animation_url)])],
                [
                    ('p', [
                        ('text', upload_heading_text),
                        ('br',),
                        ('text', 'or '),
                        ('em', upload_heading_em),
                    ]),
                    ('p', [('link', upload_cta_text, upload_cta_url)]),
                    ('p', [('text', file_restrictions_text)]),
                    ('p', [
                        ('text', 'By uploading your image or video, you agree to the Adobe '),
                        ('link', 'Terms of Use', terms_url),
                        ('text', ' and '),
                        ('link', 'Privacy Policy', privacy_url),
                    ]),
                ],
            ],
            [
                [('p', [('text', 'Quick-Action')])],
                [('p', [('text', quick_action_id)])],
            ],
        ],
        col_widths=list(col_widths),
    )


def add_frictionless_quick_action_mobile(
    doc, *, headline, subhead, tagline, upload_animation_url,
    tap_prefix_text, tap_em_text,
    file_restrictions_text, fallback_fragment_url, quick_action_id,
    terms_url=TERMS_URL, privacy_url=PRIVACY_URL,
    col_widths=(3.3, 3.3),
):
    """Mobile frictionless-quick-action-mobile hero block (5-row pattern)."""
    return add_block(
        doc, "frictionless-quick-action-mobile",
        [
            [
                [
                    ('h', 1, headline),
                    ('p', [('text', subhead)]),
                    ('p', [('text', tagline)]),
                ],
                [('p', [('text', '')])],
            ],
            [
                [('p', [('link', "Alternate video source (MP4)", upload_animation_url)])],
                [('p', [('text', tap_prefix_text), ('em', tap_em_text)])],
            ],
            [
                [('p', [('text', '')])],
                [
                    ('p', [('text', file_restrictions_text)]),
                    ('p', [
                        ('text', 'By uploading your image or video, you agree to the Adobe '),
                        ('link', 'Terms of Use', terms_url),
                        ('text', ' and '),
                        ('link', 'Privacy Policy', privacy_url),
                    ]),
                ],
            ],
            [
                [('p', [('text', 'fallback')])],
                [('p', [('link', fallback_fragment_url, fallback_fragment_url)])],
            ],
            [
                [('p', [('text', 'Quick-Action')])],
                [('p', [('text', quick_action_id)])],
            ],
        ],
        col_widths=list(col_widths),
    )


def add_how_to_steps(
    doc, *, heading, steps, variant="highlight, image, schema",
    col_widths=(2.0, 4.6),
):
    """How-to-steps block. `steps` is a list of dicts:
        { 'icon_url': ..., 'icon_alt': ..., 'title': ..., 'body': ... }
    Renders an h2 heading paragraph above the block automatically.
    """
    add_h2(doc, heading)
    rows = [
        [
            [('img', s['icon_url'], s['icon_alt'])],
            [('h', 3, s['title']), ('p', [('text', s['body'])])],
        ]
        for s in steps
    ]
    return add_block(doc, f"steps ({variant})", rows, col_widths=list(col_widths))


def add_content_column(
    doc, *, image_url, image_alt, heading, body,
    image_side='left', col_widths=(3.3, 3.3),
):
    """Single `columns` block with image + heading + body. `image_side` picks the half."""
    image_cell = [('img', image_url, image_alt)]
    text_cell = [('h', 2, heading), ('p', [('text', body)])]
    cells = [image_cell, text_cell] if image_side == 'left' else [text_cell, image_cell]
    return add_block(doc, "columns", [cells], col_widths=list(col_widths))


def add_banner(doc, *, heading, variant=None, cta=None, col_widths=(6.6,)):
    """Purple/indigo promo band. Default variant = solid `--color-info-accent` bg.

    `variant` examples: None (default), 'compact', 'standout', 'narrow', 'cool', 'light'.
    `cta` = optional (text, url) tuple rendered as a pill button below the heading.
    """
    block_name = "banner" if not variant else f"banner ({variant})"
    content = [('h', 2, heading)]
    if cta:
        content.append(('p', [('link', cta[0], cta[1])]))
    return add_block(doc, block_name, [[content]], col_widths=list(col_widths))


def add_link_list(doc, *, heading, links, col_widths=(6.6,)):
    """`link-list` block with an h3 heading and a list of (text, url) links."""
    content = [('h', 3, heading)]
    content.extend([('p', [('link', t, u)]) for t, u in links])
    return add_block(doc, "link-list", [[content]], col_widths=list(col_widths))


def add_faq(doc, *, heading, qa_pairs, col_widths=(3.3, 3.3)):
    """FAQ block. `qa_pairs` is a list of (question, answer) tuples where answer
    is either a plain string or a list of `add_runs` parts (for mixed content).
    Renders an h2 heading above the block automatically.
    """
    add_h2(doc, heading)
    rows = []
    for q, a in qa_pairs:
        if isinstance(a, str):
            a_parts = [('text', a)]
        else:
            a_parts = a
        rows.append([
            [('p', [('text', q)])],
            [('p', a_parts)],
        ])
    return add_block(doc, "faq", rows, col_widths=list(col_widths))


def add_breadcrumbs(doc, *, crumbs, col_widths=(6.6,)):
    """`breadcrumbs` block. `crumbs` = list of (text, url_or_None) tuples.
    Last crumb (current page) typically has url=None → rendered as plain text.
    """
    content = []
    for text, url in crumbs:
        if url:
            content.append(('p', [('link', text, url)]))
        else:
            content.append(('p', [('text', text)]))
    return add_block(doc, "breadcrumbs", [[content]], col_widths=list(col_widths))


def add_metadata(doc, pairs, col_widths=(3.3, 3.3)):
    """Page-level `metadata` block. `pairs` = dict (ordered in Python 3.7+)
    or list of (key, value) tuples. A value that is a URL string starting with
    http(s) is rendered as a self-referential hyperlink.
    """
    items = pairs.items() if isinstance(pairs, dict) else pairs
    rows = []
    for k, v in items:
        if isinstance(v, str) and (v.startswith('http://') or v.startswith('https://')):
            val_parts = [('link', v, v)]
        else:
            val_parts = [('text', str(v))]
        rows.append([
            [('p', [('text', k)])],
            [('p', val_parts)],
        ])
    return add_block(doc, "metadata", rows, col_widths=list(col_widths))


# ----- re-exports ----------------------------------------------------------
# Feature build.py scripts typically need: Document + add_section_break +
# the high-level block helpers. Low-level add_block/add_runs/write_cell are
# exposed for non-standard blocks the helpers don't cover.

__all__ = [
    # Core types
    'Document',
    'Pt', 'Cm', 'Inches', 'RGBColor',
    # Constants
    'HEADER_FILL', 'BORDER_COLOR', 'LINK_COLOR', 'TERMS_URL', 'PRIVACY_URL',
    # Low-level helpers (advanced use)
    'set_cell_shading', 'set_cell_borders', 'set_table_borders',
    'merge_row_cells',
    'add_hyperlink', 'add_runs', 'write_cell',
    'add_block', 'add_section_break',
    # High-level schema helpers (preferred)
    'add_h2',
    'add_showwith', 'add_section_metadata',
    'add_columns_fullsize_hero',
    'add_frictionless_quick_action',
    'add_frictionless_quick_action_mobile',
    'add_how_to_steps',
    'add_content_column',
    'add_banner',
    'add_link_list',
    'add_faq',
    'add_breadcrumbs',
    'add_metadata',
]
